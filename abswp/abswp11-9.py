import docx

d = docx.Document('c:\\users\\rimuru\\downloads\\demo.docx')
#print(d.paragraphs)

print(d.paragraphs[0].text)
print(d.paragraphs[1].text)
p = d.paragraphs[1]
print(p.runs) #for changes like bold and italics

print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)
print(p.runs[1].bold)
print(p.runs[0].bold)

print(p.runs[3].italic)

p.runs[3].underline = True
p.runs[3].text = 'italic and underlined.'
#d.save('c:\\users\\rimuru\\documents\\demo3.docx')
print(p.style)
p.style = 'Title'
#d.save('c:\\users\\rimuru\\documents\\demo2.docx')



e = docx.Document()
e.add_paragraph('Hello this is a paragraph')
e.add_paragraph('Hello this is another paragraph')
#e.save('c:\\users\\rimuru\\documents\\demo4.docx')
p = e.paragraphs[0]
print(p.add_run('This is a new run.'))
p.runs[1].bold = True
#e.save('c:\\users\\rimuru\\documents\\demo5.docx')

print('\n')

import docx

def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs: #.paragraph is the list of paragraphs in the doc
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('c:\\users\\rimuru\\downloads\\demo.docx'))


#open a new word file with docx.Document()
